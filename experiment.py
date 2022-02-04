from base64 import encode
import sys
from click import echo
from docx import Document, shared
import subprocess
import subprocess_tee
import requests
import os


class Experiment(object):
    def __init__(
        self,
        path_to_py_file: str,
        title,
        date,
        immediate: bool = True,
        manual_output=False,
        theme="seti",
    ) -> None:
        self.py_file = path_to_py_file
        self.output_file = self.py_file.replace(".py", ".txt")
        self.title = title
        self.date = date
        self.manual_output = manual_output
        self.options = {}
        self.options["theme"] = theme or "seti"
        if immediate:
            self.run()
            self.get_images()
            self.create_docx()
            self._cleanup()

    def run(self) -> None:
        command = ["python3" if sys.platform == "linux" else "python", self.py_file]
        if not self.manual_output:
            subprocess.call(
                command,
                stdout=open(self.output_file, "w"),
            )
        else:
            from asyncsubp import execute
            def cb(out: bytes):
                out = out.decode("utf-8")
                print(out)
                open(self.output_file, "w").write(out)
            execute(command, cb, lambda x: echo(x))

    def get_images(self) -> None:
        self.run()
        with open(self.py_file.replace(".py", "") + "_script.jpg", "wb") as f:
            f.write(
                requests.post(
                    "https://carbonara-42.herokuapp.com/api/cook",
                    json={
                        "code": open(self.py_file, "r").read(),
                        "lineNumbers": True,
                        "paddingHorizontal": "0px",
                        "paddingVertical": "0px",
                        "windowControls": False,
                        "theme": self.options["theme"],
                        "language": "python",
                    },
                ).content
            )
        with open(self.output_file.replace(".txt", "") + "_out.jpg", "wb") as f:
            f.write(
                requests.post(
                    "https://carbonara-42.herokuapp.com/api/cook",
                    json={
                        "code": open(self.output_file, "r").read() or " ",
                        "lineNumbers": True,
                        "paddingHorizontal": "0px",
                        "paddingVertical": "0px",
                        "windowControls": False,
                        "theme": self.options["theme"],
                        "language": "plain text",
                    },
                ).content
            )

    def create_docx(self) -> None:
        current_file = os.path.realpath(__file__)
        path_to_template = os.path.join(os.path.dirname(current_file), "template.docx")
        doc = Document(path_to_template)
        p = doc.add_paragraph("")
        r = p.add_run(f"EXPT:")
        r.bold = True
        r.font.size = shared.Pt(14)
        r.font.name = "Times New Roman"
        r.underline = True
        r = p.add_run(f"          ")
        r = p.add_run(f"{self.title.upper()}")
        r.bold = True
        r.font.size = shared.Pt(14)
        r.font.name = "Times New Roman"
        r.underline = True
        p = doc.add_paragraph("")
        r = p.add_run(f"DATE: {self.date}")
        r.bold = True
        r.font.size = shared.Pt(14)
        r.font.name = "Times New Roman"
        r.underline = True
        p = doc.add_paragraph("")
        r = p.add_run(f"CODE:")
        r.bold = True
        r.font.size = shared.Pt(14)
        r.font.name = "Times New Roman"
        r.underline = True
        doc.add_picture(self.py_file.replace(".py", "_script.jpg"), width=shared.Cm(17))
        p = doc.add_paragraph("")
        r = p.add_run(f"OUTPUT:")
        r.font.size = shared.Pt(14)
        r.font.name = "Times New Roman"
        r.underline = True
        doc.add_picture(self.py_file.replace(".py", "_out.jpg"), width=shared.Cm(17))
        doc.save(self.py_file.replace(".py", "") + ".docx")

    def _cleanup(self):
        os.remove(self.py_file.replace(".py", "_script.jpg"))
        os.remove(self.py_file.replace(".py", "_out.jpg"))
        os.remove(self.output_file)
